#!/usr/bin/env python3
"""
Advanced Markdown to DOCX Converter

This script converts Markdown files to DOCX format while preserving structure,
tables, code blocks, and applying proper styling. It also converts Mermaid diagrams
to images and embeds them in the document.

Usage:
    python md_to_docx_converter.py [input.md] [output.docx]
    python md_to_docx_converter.py  # Opens file dialog or lists MD files in current directory
"""

import os
import sys
import re
import argparse
import tempfile
import subprocess
import uuid
import shutil
from pathlib import Path
import markdown
from markdown.extensions.tables import TableExtension
from markdown.extensions.fenced_code import FencedCodeExtension
from markdown.extensions.codehilite import CodeHiliteExtension
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def setup_argparse():
    """Set up command line argument parsing."""
    parser = argparse.ArgumentParser(description='Convert Markdown to DOCX with advanced formatting')
    parser.add_argument('input_file', nargs='?', default="select", 
                        help='Input Markdown file (default: open file dialog or list files in current directory)')
    parser.add_argument('output_file', nargs='?', 
                        help='Output DOCX file (default: same name as input with .docx extension)')
    parser.add_argument('--toc', action='store_true', help='Generate table of contents')
    parser.add_argument('--style', choices=['default', 'professional', 'academic'], default='default',
                        help='Document style preset')
    parser.add_argument('--no-mermaid', action='store_true', help='Disable Mermaid diagram rendering')
    parser.add_argument('--temp-dir', help='Directory to store temporary files (default: system temp directory)')
    return parser.parse_args()


def select_file_dialog():
    """Open a file dialog to select a Markdown file."""
    try:
        import tkinter as tk
        from tkinter import filedialog
        
        root = tk.Tk()
        root.withdraw()  # Hide the main window
        
        file_path = filedialog.askopenfilename(
            title="Select Markdown File",
            filetypes=[("Markdown files", "*.md"), ("All files", "*.*")]
        )
        
        return file_path if file_path else None
    except ImportError:
        print("Tkinter not available. Please provide the file path as a command-line argument.")
        return None


def find_markdown_files(directory="."):
    """Find all Markdown files in the given directory."""
    md_files = list(Path(directory).glob("*.md"))
    return md_files


def check_mermaid_cli():
    """Check if mermaid-cli is installed."""
    try:
        # Try multiple ways to detect mermaid-cli
        commands = [
            ['npx', 'mmdc', '--version'],
            ['mmdc', '--version'],
            ['node_modules/.bin/mmdc', '--version'],
            ['./node_modules/.bin/mmdc', '--version']
        ]
        
        for cmd in commands:
            try:
                result = subprocess.run(cmd, 
                                      stdout=subprocess.PIPE, 
                                      stderr=subprocess.PIPE, 
                                      text=True)
                if result.returncode == 0:
                    print(f"Found mermaid-cli: {' '.join(cmd)}")
                    return True
            except Exception:
                continue
        
        # If we get here, none of the commands worked
        print("Warning: mermaid-cli not found. Mermaid diagrams will be included as code blocks.")
        print("To install mermaid-cli, run: npm install -g @mermaid-js/mermaid-cli")
        return False
    except Exception as e:
        print(f"Warning: Error checking for mermaid-cli: {e}")
        print("Mermaid diagrams will be included as code blocks.")
        return False


def convert_markdown_to_html(md_content):
    """Convert Markdown content to HTML."""
    # Use extensions for tables, fenced code blocks, and syntax highlighting
    extensions = [
        TableExtension(),
        FencedCodeExtension(),
        CodeHiliteExtension(noclasses=True),
        'markdown.extensions.nl2br',
        'markdown.extensions.sane_lists',
        'markdown.extensions.smarty',
        'markdown.extensions.toc',
    ]
    
    # Convert markdown to HTML
    html_content = markdown.markdown(md_content, extensions=extensions)
    return html_content


def create_docx_styles(doc, style_preset='default'):
    """Create and configure styles for the DOCX document."""
    # Heading styles
    for i in range(1, 7):
        style_name = f'Heading {i}'
        if style_name not in doc.styles:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[style_name]
        
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(20 - (i * 2))  # Decrease size for each heading level
        font.bold = True
        
        if i == 1:
            font.color.rgb = RGBColor(0, 0, 139)  # Dark blue for H1
        elif i == 2:
            font.color.rgb = RGBColor(0, 0, 180)  # Blue for H2
    
    # Code block style
    if 'Code Block' not in doc.styles:
        code_style = doc.styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        font = code_style.font
        font.name = 'Consolas'
        font.size = Pt(9)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.right_indent = Inches(0.5)
        code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Table style
    if 'Table Style' not in doc.styles:
        table_style = doc.styles.add_style('Table Style', WD_STYLE_TYPE.TABLE)
        font = table_style.font
        font.name = 'Calibri'
        font.size = Pt(10)
    
    # List styles
    if 'List Bullet' not in doc.styles:
        list_style = doc.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
        list_style.base_style = doc.styles['Normal']
        list_style.paragraph_format.left_indent = Inches(0.25)
        list_style.paragraph_format.first_line_indent = Inches(-0.25)
    
    if 'List Number' not in doc.styles:
        list_style = doc.styles.add_style('List Number', WD_STYLE_TYPE.PARAGRAPH)
        list_style.base_style = doc.styles['Normal']
        list_style.paragraph_format.left_indent = Inches(0.25)
        list_style.paragraph_format.first_line_indent = Inches(-0.25)


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a new run object (a wrapper over a run of text)
    new_run = OxmlElement('w:r')
    
    # Create a new text object
    rPr = OxmlElement('w:rPr')
    
    # Join all the xml elements together
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    # Create a new run in the paragraph
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for hyperlinks
    r.font.underline = True
    
    return hyperlink


def process_table(doc, table_elem):
    """Process an HTML table element and add it to the DOCX document."""
    rows = table_elem.find_all('tr')
    if not rows:
        return
    
    # Count columns based on the first row
    header_cells = rows[0].find_all(['th', 'td'])
    col_count = len(header_cells)
    
    # Create table in the document
    table = doc.add_table(rows=0, cols=col_count)
    table.style = 'Table Grid'
    
    # Process header row
    header_row = table.add_row()
    for i, cell in enumerate(header_cells):
        header_row.cells[i].text = cell.get_text().strip()
        # Make header bold
        for paragraph in header_row.cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Process data rows
    for row in rows[1:]:
        cells = row.find_all(['td', 'th'])
        if cells:
            row_cells = table.add_row().cells
            for i, cell in enumerate(cells):
                if i < col_count:  # Ensure we don't exceed the column count
                    row_cells[i].text = cell.get_text().strip()
    
    # Add space after table
    doc.add_paragraph()


def is_mermaid_diagram(code_text, class_attr=None):
    """Determine if a code block is a Mermaid diagram."""
    if class_attr and 'mermaid' in class_attr:
        return True
    
    # Check for common Mermaid diagram types
    mermaid_starters = [
        'graph ', 'flowchart ', 'sequenceDiagram', 'classDiagram',
        'stateDiagram', 'erDiagram', 'journey', 'gantt', 'pie',
        'gitGraph', 'requirementDiagram'
    ]
    
    code_start = code_text.strip().split('\n')[0].strip()
    return any(code_start.startswith(starter) for starter in mermaid_starters)


def render_mermaid_diagram(code_text, output_dir, diagram_id=None):
    """Render a Mermaid diagram to an image file."""
    if diagram_id is None:
        diagram_id = str(uuid.uuid4())
    
    # Create temporary mermaid file
    mermaid_file = os.path.join(output_dir, f"diagram_{diagram_id}.mmd")
    output_file = os.path.join(output_dir, f"diagram_{diagram_id}.png")
    
    with open(mermaid_file, 'w', encoding='utf-8') as f:
        f.write(code_text)
    
    # Try multiple commands to render the diagram
    commands = [
        ['npx', 'mmdc', '-i', mermaid_file, '-o', output_file, '-b', 'transparent'],
        ['mmdc', '-i', mermaid_file, '-o', output_file, '-b', 'transparent'],
        ['node_modules/.bin/mmdc', '-i', mermaid_file, '-o', output_file, '-b', 'transparent'],
        ['./node_modules/.bin/mmdc', '-i', mermaid_file, '-o', output_file, '-b', 'transparent']
    ]
    
    for cmd in commands:
        try:
            print(f"Trying to render with: {' '.join(cmd)}")
            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True
            )
            
            if result.returncode == 0 and os.path.exists(output_file):
                print(f"Successfully rendered diagram to {output_file}")
                return output_file
        except Exception as e:
            print(f"Command failed: {e}")
            continue
    
    print("All rendering attempts failed. Including diagram as code block.")
    return None


def add_image_to_docx(doc, image_path, width=None):
    """Add an image to the DOCX document."""
    if not os.path.exists(image_path):
        return
    
    if width is None:
        width = Inches(6)  # Default width
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)
    
    # Add space after image
    doc.add_paragraph()


def process_code_block(doc, pre_elem, temp_dir, render_mermaid=True):
    """Process a code block and add it to the document with proper formatting."""
    code_text = pre_elem.get_text()
    
    # Check if it's a Mermaid diagram
    if is_mermaid_diagram(code_text, pre_elem.get('class', [])):
        if render_mermaid:
            # Try to render the diagram as an image
            image_path = render_mermaid_diagram(code_text, temp_dir)
            if image_path:
                # Add a caption
                p = doc.add_paragraph()
                p.add_run("Mermaid Diagram:").bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add the rendered image
                add_image_to_docx(doc, image_path)
                return
        
        # Fallback if rendering fails or is disabled
        p = doc.add_paragraph()
        p.add_run("Mermaid Diagram:").bold = True
        p = doc.add_paragraph(code_text, style='Code Block')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()  # Add space after diagram
    else:
        # Regular code block
        p = doc.add_paragraph(code_text, style='Code Block')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add a light gray shading to the paragraph
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F5F5F5')  # Light gray
        p._p.get_or_add_pPr().append(shading_elm)
        
        doc.add_paragraph()  # Add space after code block


def process_list(doc, list_elem, list_type):
    """Process an HTML list element and add it to the DOCX document."""
    items = list_elem.find_all('li', recursive=False)
    
    for i, item in enumerate(items):
        # Determine the list style and prefix
        if list_type == 'ol':
            p = doc.add_paragraph(style='List Number')
            prefix = f"{i+1}. "
        else:  # ul
            p = doc.add_paragraph(style='List Bullet')
            prefix = "• "
        
        # Add the list item text with proper prefix
        p.add_run(prefix + item.get_text().strip())
        
        # Handle nested lists
        nested_lists = item.find_all(['ul', 'ol'], recursive=False)
        for nested_list in nested_lists:
            process_list(doc, nested_list, nested_list.name)


def html_to_docx(html_content, output_path, temp_dir, render_mermaid=True):
    """Convert HTML content to a DOCX file."""
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create a new DOCX document
    doc = Document()
    
    # Set up document styles
    create_docx_styles(doc)
    
    # Process HTML elements
    for elem in soup.find_all(recursive=False):
        if elem.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Handle headings
            level = int(elem.name[1])
            p = doc.add_paragraph(elem.get_text().strip(), style=f'Heading {level}')
        
        elif elem.name == 'p':
            # Handle paragraphs
            p = doc.add_paragraph()
            
            # Process paragraph content, preserving links
            for content in elem.contents:
                if content.name == 'a':
                    add_hyperlink(p, content.get_text(), content.get('href', '#'))
                elif content.name == 'strong' or content.name == 'b':
                    p.add_run(content.get_text()).bold = True
                elif content.name == 'em' or content.name == 'i':
                    p.add_run(content.get_text()).italic = True
                elif content.name == 'code':
                    run = p.add_run(content.get_text())
                    run.font.name = 'Consolas'
                else:
                    if hasattr(content, 'get_text'):
                        p.add_run(content.get_text())
                    else:
                        p.add_run(str(content))
        
        elif elem.name == 'table':
            # Handle tables
            process_table(doc, elem)
        
        elif elem.name == 'pre':
            # Handle code blocks
            process_code_block(doc, elem, temp_dir, render_mermaid)
        
        elif elem.name == 'ul':
            # Handle unordered lists
            process_list(doc, elem, 'ul')
        
        elif elem.name == 'ol':
            # Handle ordered lists
            process_list(doc, elem, 'ol')
        
        elif elem.name == 'hr':
            # Handle horizontal rules
            doc.add_paragraph('_' * 50)
        
        elif elem.name == 'blockquote':
            # Handle blockquotes
            p = doc.add_paragraph(elem.get_text().strip())
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            # Add a light gray left border
            p.paragraph_format.left_indent = Inches(0.5)
    
    # Save the document
    doc.save(output_path)
    return output_path


def process_nested_elements(doc, element, level=0):
    """Process nested HTML elements recursively."""
    if hasattr(element, 'contents'):
        for child in element.contents:
            if hasattr(child, 'name'):
                if child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # Handle headings
                    heading_level = int(child.name[1])
                    p = doc.add_paragraph(child.get_text().strip(), style=f'Heading {heading_level}')
                
                elif child.name == 'p':
                    # Handle paragraphs
                    p = doc.add_paragraph(child.get_text().strip())
                
                elif child.name == 'table':
                    # Handle tables
                    process_table(doc, child)
                
                elif child.name == 'pre':
                    # Handle code blocks
                    process_code_block(doc, child)
                
                elif child.name == 'ul':
                    # Handle unordered lists
                    process_list(doc, child, 'ul')
                
                elif child.name == 'ol':
                    # Handle ordered lists
                    process_list(doc, child, 'ol')
                
                # Recursively process nested elements
                process_nested_elements(doc, child, level + 1)


def get_input_file():
    """Get the input file path through various methods."""
    # First, try to use file dialog
    file_path = select_file_dialog()
    if file_path:
        return Path(file_path)
    
    # If dialog fails or is cancelled, look for MD files in current directory
    md_files = find_markdown_files()
    if not md_files:
        print("No Markdown files found in the current directory.")
        print("Please provide a file path: python md_to_docx_converter.py <path_to_markdown_file>")
        sys.exit(1)
    elif len(md_files) == 1:
        input_path = md_files[0]
        print(f"Found one Markdown file: {input_path}")
        return input_path
    else:
        print("Multiple Markdown files found. Please select one:")
        for i, file in enumerate(md_files):
            print(f"{i+1}. {file}")
        
        try:
            selection = int(input("Enter the number of the file to convert: "))
            if 1 <= selection <= len(md_files):
                return md_files[selection-1]
            else:
                print("Invalid selection. Exiting.")
                sys.exit(1)
        except ValueError:
            print("Invalid input. Exiting.")
            sys.exit(1)


def main():
    """Main function to convert Markdown to DOCX."""
    args = setup_argparse()
    
    # Get input file path
    if args.input_file == "select":
        input_path = get_input_file()
    else:
        input_path = Path(args.input_file)
        if not input_path.exists():
            print(f"Error: Input file '{input_path}' not found.")
            print("Trying to find Markdown files in the current directory...")
            input_path = get_input_file()
    
    # Determine output path
    if args.output_file:
        output_path = Path(args.output_file)
    else:
        output_path = input_path.with_suffix('.docx')
    
    # Check for mermaid-cli if not disabled
    render_mermaid = False
    if not args.no_mermaid:
        render_mermaid = check_mermaid_cli()
    
    # Create temporary directory for processing
    temp_dir = args.temp_dir if args.temp_dir else tempfile.mkdtemp()
    try:
        os.makedirs(temp_dir, exist_ok=True)
        
        print(f"Converting '{input_path}' to '{output_path}'...")
        
        # Read Markdown content
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert Markdown to HTML
        html_content = convert_markdown_to_html(md_content)
        
        # Convert HTML to DOCX
        output_file = html_to_docx(html_content, output_path, temp_dir, render_mermaid)
        
        print(f"Conversion complete! Output saved to '{output_file}'")
    
    finally:
        # Clean up temporary files if we created them
        if not args.temp_dir and os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)


if __name__ == "__main__":
    main()
